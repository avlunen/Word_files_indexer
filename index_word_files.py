#-------------------------------------------------------------------------------
# Name:        Indexer
# Purpose:     To extract a list of words from MS Word documents in order to
#              help creating an index, v2.0
#
# Author:      Alexander von Lunen
#
# Created:     04/08/2015
# Copyright:   (c) Alexander von Lunen 2015
# Licence:     Public Domain
#-------------------------------------------------------------------------------
import glob
import os
import time
import io
from topia.termextract import tag
from topia.termextract import extract
from docx import *

# write a list of index words into the first column of a two-column table
# in a Word document
def write_concordance(in_list, file_name):
   doc = Document()

   table = doc.add_table(rows=1, cols=1)

   for row in in_list:
      row_cells = table.add_row().cells
      row_cells[0].text = row

   doc.save(file_name)
   del doc

# Retrieve the text of a Word document from a docx object
def getdocumenttext(document):
   return '\n'.join([paragraph.text
          for paragraph in document.paragraphs])

# main routine
def main():
   try:
       # list of index terms
       index_list = list()

       # init tagging
       tagger = tag.Tagger()
       tagger.initialize()
       extractor = extract.TermExtractor(tagger)
       #extractor.filter = extract.permissiveFilter
       #extractor.filter = extract.DefaultFilter(singleStrengthMinOccur=2)

       # get file path
       p = os.path.join('final.ms'+ os.sep, '*chapter*.docx') # you may need to customize this

       # go through files
       for infile in glob.glob(p):
          # open document
          doc = Document(os.getcwd()+'\\'+infile)
          print os.getcwd()+'\\'+infile

          # get text from Word document
          text = getdocumenttext(doc)

          # tagging
          l = extractor(text)
          for item in l:
             if item[0] not in index_list:
                index_list.append(item[0])

          # close Word document
          del doc

       write_concordance(sorted(index_list), os.getcwd()+os.sep+'all_concordance.docx')
   finally:
      print "Done!"

if __name__ == '__main__':
    main()
